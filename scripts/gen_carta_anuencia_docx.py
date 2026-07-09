#!/usr/bin/env python3
"""
Gera o modelo de CARTA DE ANUÊNCIA exigida no item 4.5 do Edital PDJ/Fiocruz
VPPCB 2026 (anuência da chefia do laboratório/departamento/líder do grupo).

O edital NÃO fornece modelo oficial de anuência; este é um modelo formal,
em uma página, com campos entre [ ] para preenchimento e assinatura.

Saída: submissao_pdj/Carta_de_Anuencia_PDJ.docx (+ PDF via LibreOffice)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submissao_pdj" / "Carta_de_Anuencia_PDJ.docx"

AZUL = RGBColor(0x1F, 0x39, 0x64)
CINZA = RGBColor(0x55, 0x55, 0x55)


def para(doc, text, size=11, justify=True, italic=False, color=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    return p


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(3.0)

    # Cabeçalho institucional
    for txt, sz, bold, color in [
        ("FUNDAÇÃO OSWALDO CRUZ — FIOCRUZ", 12, True, AZUL),
        ("Instituto Aggeu Magalhães (IAM) — Fiocruz Pernambuco", 11, False, CINZA),
        ("[Laboratório / Departamento / Grupo de Pesquisa]", 10, False, CINZA),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.font.bold = bold
        r.font.color.rgb = color

    doc.add_paragraph()

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CARTA DE ANUÊNCIA")
    r.font.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = AZUL
    doc.add_paragraph()

    # Local e data
    para(doc, "Recife, PE, ______ de ______________________ de 2026.",
         justify=False, space_after=14)

    # Destinatário
    para(doc, "À Coordenação de Fomento à Pesquisa", justify=False, space_after=0)
    para(doc, "Vice-Presidência de Pesquisa e Coleções Biológicas (VPPCB) — Fiocruz",
         justify=False, space_after=0)
    para(doc, "Programa de Pós-Doutorado Júnior (PDJ) — Edital 2026",
         justify=False, space_after=14)

    # Assunto
    p = doc.add_paragraph()
    r = p.add_run("Assunto: ")
    r.font.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run("Anuência para realização de subprojeto de Pós-Doutorado Júnior (PDJ) "
                   "no âmbito do laboratório/grupo de pesquisa.")
    r2.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(14)

    # Corpo
    para(doc, "Prezados(as) Senhores(as),", justify=False, space_after=10)

    para(doc,
         "Na qualidade de [cargo — ex.: chefe do Laboratório / chefe do Departamento / "
         "líder do Grupo de Pesquisa] do(a) [nome do laboratório / departamento / grupo] "
         "do Instituto Aggeu Magalhães (IAM) — Fiocruz Pernambuco, declaro estar CIENTE e "
         "DE ACORDO com a realização do subprojeto intitulado \u201cJEPA-Spillover: "
         "aprendizado preditivo em espaço latente para vigilância genômica de vírus com "
         "potencial zoonótico\u201d, a ser desenvolvido pelo(a) candidato(a) Gabriel Bezerra "
         "Motta Câmara, sob supervisão do(a) Prof.(a) [nome do(a) supervisor(a)], no âmbito "
         "do Edital de Pós-Doutorado Júnior (PDJ) da VPPCB/Fiocruz, vigência 2026–2027.",
         space_after=10)

    para(doc,
         "Declaro, ainda, que o(a) referido(a) laboratório/grupo de pesquisa dispõe da "
         "infraestrutura necessária — instalações, recursos computacionais e acesso a bases "
         "de dados — para o adequado desenvolvimento das atividades previstas no subprojeto, "
         "e que me comprometo a assegurar ao(à) bolsista as condições necessárias à sua "
         "execução durante toda a vigência da bolsa.",
         space_after=10)

    para(doc,
         "Manifesto, por meio desta, minha anuência para que o subprojeto seja conduzido "
         "nas dependências desta Unidade, em articulação com o projeto de pesquisa do(a) "
         "supervisor(a) ao qual se vincula.",
         space_after=18)

    para(doc, "Atenciosamente,", justify=False, space_after=28)

    # Assinatura
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("_______________________________________________")
    r.font.size = Pt(11)
    for txt in [
        "[Nome completo da chefia / líder do grupo]",
        "[Cargo / função]",
        "[Laboratório / Departamento / Grupo de Pesquisa]",
        "Instituto Aggeu Magalhães — Fiocruz Pernambuco",
        "[Carimbo institucional, se aplicável]",
    ]:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = q.add_run(txt)
        rr.font.size = Pt(10)
        rr.font.color.rgb = CINZA if txt.startswith("[") else RGBColor(0, 0, 0)
        q.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    nota = doc.add_paragraph()
    rn = nota.add_run(
        "Observação: preencher os campos entre colchetes [ ]. Recomenda-se imprimir em papel "
        "timbrado do laboratório/IAM e assinar. O edital (item 4.5) exige a anuência da chefia "
        "do laboratório, do departamento ou do líder do grupo de pesquisa da Unidade onde o "
        "subprojeto será realizado."
    )
    rn.italic = True
    rn.font.size = Pt(8.5)
    rn.font.color.rgb = CINZA

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
