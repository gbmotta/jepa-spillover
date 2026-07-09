#!/usr/bin/env python3
"""
Gera o CHECKLIST de submissão (Edital PDJ/Fiocruz VPPCB 2026) em .docx,
com caixas de marcação, para acompanhamento e compartilhamento.

Saída: submissao_pdj/Checklist_Submissao_PDJ.docx (+ PDF via LibreOffice)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submissao_pdj" / "Checklist_Submissao_PDJ.docx"

AZUL = RGBColor(0x1F, 0x39, 0x64)
CINZA = RGBColor(0x55, 0x55, 0x55)
VERDE = RGBColor(0x1B, 0x5E, 0x20)


def set_cell_bg(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcpr.append(shd)


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(13)
    return p


def para(doc, text, size=11, justify=False, italic=False, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    return p


def check(doc, text, done=False, tag=None):
    """Linha de checklist com caixa."""
    p = doc.add_paragraph()
    box = p.add_run(("\u2611  " if done else "\u2610  "))
    box.font.size = Pt(12)
    if done:
        box.font.color.rgb = VERDE
    if tag:
        rt = p.add_run(f"[{tag}] ")
        rt.font.size = Pt(9)
        rt.font.bold = True
        rt.font.color.rgb = CINZA
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def make_table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, "1F3964")
        run = c.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Checklist de Submissão — Pós-Doutorado Júnior (PDJ)")
    r.font.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = AZUL
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Fiocruz — VPPCB · Instituto Aggeu Magalhães · Edital 2026 (Bolsa Nova)")
    r2.font.size = Pt(10)
    r2.font.color.rgb = CINZA

    # Info
    for txt in [
        "Candidato: Gabriel Bezerra Motta Câmara   |   Supervisor: [preencher]",
        "Submissão: sistema Fomento à Pesquisa on-line (fomentoapesquisa.fiocruz.br)",
        "Prazo de inscrição: 9 de junho a 10 de julho de 2026",
    ]:
        para(doc, txt, size=10, color=CINZA)

    para(doc, "Legenda: SUP = depende do supervisor · CAND = depende do candidato · PRONTO = documento já preparado por nós",
         size=8.5, italic=True, color=CINZA)

    # 1. Supervisor
    h1(doc, "1. Documentos do Supervisor (Túlio)")
    check(doc, "Termo de inscrição on-line (2026–2027) — assinado por supervisor E bolsista (gerado no sistema)", tag="SUP")
    check(doc, "Projeto de pesquisa do orientador (item 4.4.b) — esqueleto pronto: Projeto_Orientador_ESQUELETO.docx → preencher", tag="SUP")
    check(doc, "CV Lattes/CNPq do supervisor — atualizado até 10/jul/2026", tag="SUP")
    check(doc, "Súmula da produção técnico-científica do supervisor (modelo no sistema)", tag="SUP")

    # 2. Candidato
    h1(doc, "2. Documentos do Candidato (Gabriel)")
    check(doc, "Termo de Consentimento Livre e Esclarecido (disponibilizado pelo orientador no sistema)", tag="CAND")
    check(doc, "CV Lattes/CNPq do candidato — atualizado até 10/jul/2026", tag="CAND")
    check(doc, "Súmula da produção técnico-científica do candidato (modelo no sistema)", tag="CAND")
    check(doc, "Comprovante de doutorado — certificado, ou ata da defesa, ou carta do orientador (defesa até 30/ago/2026)", tag="CAND")
    check(doc, "Subprojeto do candidato (item 4.4.f) — Subprojeto_JEPA-Spillover_PDJ.docx (falta preencher campos [ ])",
          done=True, tag="PRONTO")

    # 3. Comum
    h1(doc, "3. Documento comum (obrigatório)")
    check(doc, "Carta de anuência da chefia (item 4.5) — modelo pronto: Carta_de_Anuencia_PDJ.docx → assinar", tag="SUP")

    # 4. Autorizações
    h1(doc, "4. Autorizações legais (item 4.6 — somente quando aplicável)")
    for t in ["Comitê de Ética em Pesquisa (CEP)", "Comitê de Ética no Uso de Animais (CEUA)",
              "SISBIO", "Comitê Interno de Biossegurança (CI-BIO)", "SISGEN"]:
        check(doc, t)
    para(doc,
         "Observação: o subprojeto JEPA-Spillover usa apenas dados públicos (sem material biológico "
         "novo, sem pacientes) → em princípio não exige essas autorizações. Confirmar com o Túlio "
         "conforme o escopo do projeto maior.",
         size=9, italic=True, color=CINZA)

    # 5. Ações afirmativas
    h1(doc, "5. Ações afirmativas (Anexo I — se o candidato optar)")
    check(doc, "Formulário de autodeclaração (negros/pardos, indígenas, pessoas trans, PcD)")
    check(doc, "Laudo médico com CID-10 (apenas para PcD)")

    # Cronograma
    h1(doc, "Prazos-chave do edital")
    make_table(
        doc,
        ["Etapa", "Data"],
        [
            ["Período de inscrição on-line", "9 jun – 10 jul 2026"],
            ["Homologação dos processos", "13–15 jul 2026"],
            ["Resultado da homologação", "16 jul 2026"],
            ["Recurso da homologação", "17–18 jul 2026"],
            ["Resultado da avaliação (bolsa nova)", "a partir de 10 ago 2026"],
            ["Recurso da avaliação", "11–12 ago 2026"],
            ["Comissão de Heteroidentificação (se cotista)", "18–19 ago 2026"],
            ["Defesa de doutorado (prazo máximo, se em conclusão)", "30 ago 2026"],
        ],
        widths=[11.0, 5.5],
        font_size=9.5,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
