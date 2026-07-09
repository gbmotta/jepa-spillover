#!/usr/bin/env python3
"""
Gera um ESQUELETO/RASCUNHO do "Projeto de pesquisa do orientador" (item 4.4.b
do Edital PDJ/Fiocruz VPPCB 2026), ao qual o subprojeto do candidato se vincula.

Estruturado segundo os critérios de avaliação do edital (Anexo IV — Bloco
Projeto): relevância científica, relevância para o SUS, desenho/abordagem
metodológica. Campos entre [ ] devem ser preenchidos pelo(a) supervisor(a).

Saída: submissao_pdj/Projeto_Orientador_ESQUELETO.docx (+ PDF via LibreOffice)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submissao_pdj" / "Projeto_Orientador_ESQUELETO.docx"

AZUL = RGBColor(0x1F, 0x39, 0x64)
CINZA = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcpr.append(shd)


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(14)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(12)
    return p


def para(doc, text, size=11, justify=True, italic=False, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    return p


def guide(doc, text):
    """Texto-guia em itálico cinza para o supervisor preencher."""
    return para(doc, text, size=10, italic=True, color=CINZA)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def make_table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, "1F3964")
        run = c.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # Capa
    for txt, sz, bold, color, align in [
        ("FUNDAÇÃO OSWALDO CRUZ — FIOCRUZ", 12, True, AZUL, "c"),
        ("Vice-Presidência de Pesquisa e Coleções Biológicas (VPPCB)", 11, False, CINZA, "c"),
        ("Instituto Aggeu Magalhães — Fiocruz Pernambuco", 11, False, CINZA, "c"),
        ("Edital de Pós-Doutorado Júnior (PDJ) — 2026", 11, False, CINZA, "c"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.font.bold = bold
        r.font.color.rgb = color

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJETO DE PESQUISA DO(A) SUPERVISOR(A)")
    r.font.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[Título do projeto de pesquisa / desenvolvimento tecnológico]")
    r.font.bold = True
    r.font.size = Pt(14)
    r.italic = True
    r.font.color.rgb = CINZA

    doc.add_paragraph()
    doc.add_paragraph()

    ident = [
        ("Coordenador(a) / Supervisor(a)", "[Nome — servidor(a) ativo(a) da Fiocruz, com título de doutor(a)]"),
        ("Unidade", "Instituto Aggeu Magalhães (IAM) — Fiocruz Pernambuco"),
        ("Laboratório / Grupo", "[Nome do laboratório ou grupo de pesquisa]"),
        ("Grande área / Linha", "[Ex.: Bioinformática / Vigilância genômica / Doenças infecciosas]"),
        ("Vínculo do subprojeto PDJ", "JEPA-Spillover (candidato: Gabriel Bezerra Motta Câmara)"),
        ("Vigência", "[Ex.: 2026–2027 / 24–36 meses]"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in ident:
        cells = t.add_row().cells
        rk = cells[0].paragraphs[0].add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)
        set_cell_bg(cells[0], "EDF1F7")
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(10)
        cells[0].width = Cm(5.5)
        cells[1].width = Cm(11.0)

    doc.add_paragraph()
    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = nota.add_run(
        "ESQUELETO PARA PREENCHIMENTO PELO(A) SUPERVISOR(A). Os textos em itálico cinza são "
        "orientações e devem ser substituídos/removidos. Campos entre [ ] devem ser preenchidos."
    )
    rn.italic = True
    rn.font.size = Pt(8.5)
    rn.font.color.rgb = CINZA

    doc.add_page_break()

    # 1. Resumo
    h1(doc, "1. Resumo")
    guide(doc,
          "Descreva em 1 parágrafo o problema central do projeto, a abordagem geral e o "
          "impacto esperado para a ciência e para o SUS. Deixe claro que o subprojeto "
          "JEPA-Spillover é um dos componentes computacionais deste projeto maior.")
    para(doc, "[Resumo do projeto...]")

    # 2. Introdução e justificativa
    h1(doc, "2. Introdução e justificativa")
    h2(doc, "2.1. Relevância científica")
    guide(doc,
          "Contextualize a área (ex.: vigilância genômica de patógenos, bioinformática, "
          "doenças infecciosas emergentes). Fundamente com literatura o problema que o "
          "projeto ataca e a lacuna de conhecimento. Critério avaliado: 'Relevância para a "
          "pesquisa científica e/ou tecnológica'.")
    para(doc, "[Texto...]")
    h2(doc, "2.2. Relevância para o SUS e para a sociedade")
    guide(doc,
          "Explique como o projeto contribui para o Sistema Único de Saúde, para políticas "
          "de saúde pública e para populações vulneráveis. Critério avaliado: 'Relevância "
          "para o SUS e para a sociedade'. O JEPA-Spillover contribui aqui ao antecipar "
          "riscos zoonóticos e complementar sistemas como InfoDengue/InfoGripe.")
    para(doc, "[Texto...]")

    # 3. Objetivos
    h1(doc, "3. Objetivos")
    h2(doc, "3.1. Objetivo geral")
    para(doc, "[Objetivo geral do projeto do orientador...]")
    h2(doc, "3.2. Objetivos específicos")
    guide(doc, "Liste os objetivos específicos. Marque com (*) aquele(s) atendido(s) pelo subprojeto JEPA-Spillover.")
    for i in range(1, 4):
        bullet(doc, f"[Objetivo específico {i}...]")
    bullet(doc, "[Objetivo específico atendido pelo subprojeto JEPA-Spillover — ex.: desenvolver "
                "métodos computacionais de priorização de vírus com potencial zoonótico].", bold_prefix="(*) ")

    # 4. Abordagem metodológica
    h1(doc, "4. Desenho experimental e abordagem teórico-metodológica")
    guide(doc,
          "Descreva materiais, métodos, desenho experimental e/ou abordagem teórica. "
          "Critério avaliado: 'Desenho experimental ou abordagem teórico-metodológica'. "
          "Inclua a componente computacional/bioinformática onde o subprojeto se encaixa.")
    para(doc, "[Metodologia...]")

    # 5. Articulação com o subprojeto
    h1(doc, "5. Articulação com o subprojeto JEPA-Spillover")
    para(doc,
         "O subprojeto de Pós-Doutorado Júnior \u201cJEPA-Spillover\u201d integra este projeto "
         "como sua componente de inteligência computacional. Enquanto o projeto do(a) "
         "supervisor(a) [descrever o eixo — ex.: vigilância genômica / epidemiologia molecular "
         "de patógenos], o subprojeto aporta métodos de aprendizado de máquina de última "
         "geração (redes JEPA / Transformers) para aprender representações latentes de genomas "
         "virais e priorizar vírus com potencial de spillover. Essa articulação agrega "
         "capacidade analítica reprodutível e escalável ao grupo, com potencial de publicações "
         "conjuntas e captação de novos financiamentos.")
    guide(doc, "Ajuste este parágrafo conforme o escopo real do seu projeto e a aderência do subprojeto.")

    # 6. Resultados esperados e impacto
    h1(doc, "6. Resultados esperados e impacto")
    guide(doc, "Liste produtos esperados (publicações, modelos, protocolos, formação de RH) e o impacto para o SUS.")
    for _ in range(3):
        bullet(doc, "[Resultado esperado...]")

    # 7. Infraestrutura
    h1(doc, "7. Infraestrutura disponível")
    guide(doc,
          "Descreva a infraestrutura do laboratório/grupo/IAM: instalações, equipamentos, "
          "recursos computacionais (GPU/HPC), bases de dados e apoio institucional. Critério "
          "avaliado (no subprojeto): 'Infraestrutura para a realização da pesquisa'.")
    para(doc, "[Infraestrutura...]")

    # 8. Cronograma
    h1(doc, "8. Cronograma")
    guide(doc, "Preencha as macro-etapas do projeto. Exequibilidade é critério de avaliação.")
    make_table(
        doc,
        ["Etapa / Meta", "Atividades", "Período"],
        [
            ["[Meta 1]", "[Atividades...]", "[Meses ...]"],
            ["[Meta 2]", "[Atividades...]", "[Meses ...]"],
            ["[Meta 3]", "[Atividades...]", "[Meses ...]"],
            ["[Meta 4]", "[Atividades...]", "[Meses ...]"],
        ],
        widths=[4.5, 8.0, 4.0],
    )

    # 9. Equipe
    h1(doc, "9. Equipe e colaborações")
    guide(doc, "Liste membros da equipe, colaboradores e redes (institucionais, nacionais, internacionais).")
    for _ in range(2):
        bullet(doc, "[Nome — função no projeto]")

    # 10. Referências
    h1(doc, "10. Referências")
    guide(doc, "Insira as referências bibliográficas do projeto.")
    para(doc, "[1] [Referência...]", size=9.5)

    # Rodapé
    footer = sec.footer
    fp = footer.paragraphs[0]
    fr = fp.add_run("Projeto do(a) Supervisor(a) — PDJ · Instituto Aggeu Magalhães / Fiocruz PE · 2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = CINZA
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
