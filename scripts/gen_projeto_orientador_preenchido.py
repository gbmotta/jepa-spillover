#!/usr/bin/env python3
"""
Gera uma versão PREENCHIDA e realista do "Projeto de pesquisa do orientador"
(item 4.4.b do Edital PDJ/Fiocruz VPPCB 2026), redigido como um projeto
guarda-chuva de bioinformática/vigilância genômica no IAM-Fiocruz, do qual o
subprojeto JEPA-Spillover é a componente de Inteligência Artificial.

Projeto fictício, porém plausível, com base no escopo real do candidato.
Saída: submissao_pdj/Projeto_Orientador_PDJ.docx (+ PDF via LibreOffice)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submissao_pdj" / "Projeto_Orientador_PDJ.docx"

AZUL = RGBColor(0x1F, 0x39, 0x64)
CINZA = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcpr.append(shd)


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(14)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(12)
    return p


def para(doc, text, size=11, justify=True, italic=False, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def make_table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, "1F3964")
        run = c.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # ---------------- Capa ----------------
    for txt, sz, bold, color in [
        ("FUNDAÇÃO OSWALDO CRUZ — FIOCRUZ", 12, True, AZUL),
        ("Vice-Presidência de Pesquisa e Coleções Biológicas (VPPCB)", 11, False, CINZA),
        ("Instituto Aggeu Magalhães — Fiocruz Pernambuco", 11, False, CINZA),
        ("Projeto de Pesquisa — vinculado ao Edital de Pós-Doutorado Júnior (PDJ) 2026", 11, False, CINZA),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.font.bold = bold
        r.font.color.rgb = color

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJETO DE PESQUISA DO(A) SUPERVISOR(A)")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "GenVig-IA — Vigilância Genômica Integrada e Inteligência Computacional para "
        "Caracterização Molecular e Predição de Risco de Vírus Emergentes e Zoonóticos "
        "de Importância em Saúde Pública no Brasil"
    )
    r.font.bold = True
    r.font.size = Pt(14)
    r.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    ident = [
        ("Coordenador(a) / Supervisor(a)", "Túlio Campos [confirmar nome completo e titulação]"),
        ("Unidade", "Instituto Aggeu Magalhães (IAM) — Fiocruz Pernambuco"),
        ("Laboratório / Núcleo", "Núcleo de Bioinformática (NBI) [confirmar]"),
        ("Grande área", "Ciências Biológicas / Bioinformática / Saúde Coletiva"),
        ("Linha de pesquisa", "Vigilância genômica, epidemiologia molecular e IA aplicada à saúde"),
        ("Subprojeto PDJ vinculado", "JEPA-Spillover — candidato Gabriel Bezerra Motta Câmara"),
        ("Duração", "36 meses (subprojeto PDJ: 12 meses, renovável)"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in ident:
        cells = t.add_row().cells
        rk = cells[0].paragraphs[0].add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)
        set_cell_bg(cells[0], "EDF1F7")
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(10)
        cells[0].width = Cm(5.5)
        cells[1].width = Cm(11.0)

    doc.add_paragraph()
    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = nota.add_run(
        "Documento de trabalho, redigido como projeto guarda-chuva plausível para "
        "contextualizar o subprojeto de PDJ. Ajustar nomes, titulação e dados institucionais "
        "(campos entre [ ]) antes da submissão."
    )
    rn.italic = True
    rn.font.size = Pt(8.5)
    rn.font.color.rgb = CINZA

    doc.add_page_break()

    # ---------------- 1. Resumo ----------------
    h1(doc, "1. Resumo")
    para(doc,
         "As doenças infecciosas emergentes e reemergentes — em especial as de origem viral e "
         "potencial zoonótico — representam uma ameaça crescente à saúde pública global e "
         "brasileira, como evidenciado pelas epidemias de Zika e chikungunya no Nordeste, pela "
         "emergência do Oropouche e pela pandemia de COVID-19. O barateamento do sequenciamento "
         "de nova geração (NGS) e a expansão das redes de vigilância genômica produziram um "
         "volume sem precedentes de genomas virais, mas a capacidade de análise, integração e, "
         "sobretudo, de predição de risco não acompanhou esse crescimento. Este projeto propõe "
         "uma plataforma integrada de bioinformática e inteligência computacional para "
         "caracterização molecular, monitoramento filodinâmico e predição de risco de vírus de "
         "importância em saúde pública, articulando dados genômicos, de hospedeiros e ecológicos. "
         "O projeto organiza-se em quatro eixos — (i) aquisição e curadoria de dados, "
         "(ii) caracterização molecular e filogenética/filodinâmica, (iii) modelagem "
         "computacional e aprendizado de máquina para predição de risco zoonótico e "
         "(iv) integração com a vigilância e disseminação. O eixo (iii) é desenvolvido, no "
         "âmbito do Pós-Doutorado Júnior, pelo subprojeto JEPA-Spillover, que aplica redes de "
         "arquitetura preditiva em espaço latente (JEPA) para priorizar vírus com maior potencial "
         "de transbordamento antes da ocorrência de surtos. Espera-se entregar uma "
         "infraestrutura analítica aberta e reprodutível, com impacto direto para a capacidade "
         "de resposta do Sistema Único de Saúde (SUS).")

    # ---------------- 2. Introdução e justificativa ----------------
    h1(doc, "2. Introdução e justificativa")
    h2(doc, "2.1. Relevância científica")
    para(doc,
         "Estima-se que a maioria das doenças infecciosas emergentes em humanos tenha origem "
         "zoonótica, e que exista um vasto 'espaço viral' ainda não caracterizado em reservatórios "
         "animais. A vigilância genômica tornou-se ferramenta central para detectar e caracterizar "
         "esses agentes, mas enfrenta dois gargalos científicos. Primeiro, a análise "
         "filogenética e filodinâmica clássica descreve o passado da transmissão, porém tem poder "
         "limitado para antecipar quais vírus, entre os milhares recém-descritos por metagenômica, "
         "têm maior probabilidade de infectar humanos. Segundo, os métodos supervisionados de "
         "predição de risco dependem de rótulos de zoonose escassos e enviesados — a ausência de "
         "registro de infecção humana frequentemente reflete falta de estudo, e não ausência de "
         "risco. Há, portanto, uma lacuna metodológica clara: são necessárias abordagens capazes "
         "de aprender, de forma pouco supervisionada, representações funcionais dos genomas virais "
         "que capturem compatibilidade com hospedeiros. Avanços recentes em aprendizado "
         "auto-supervisionado — em particular as arquiteturas preditivas em espaço latente (JEPA) "
         "— abrem uma via promissora e ainda inexplorada nesse domínio.")
    h2(doc, "2.2. Relevância para o SUS e para a sociedade")
    para(doc,
         "O Brasil convive com uma carga elevada e recorrente de arboviroses (dengue, Zika, "
         "chikungunya, Oropouche, febre amarela) e com o risco permanente de introdução e "
         "emergência de novos patógenos. O Nordeste, região de atuação do Instituto Aggeu "
         "Magalhães, foi o epicentro da epidemia de Zika e da síndrome congênita associada, "
         "expondo a necessidade de capacidade local de resposta rápida. Este projeto contribui "
         "diretamente para o SUS ao fortalecer a vigilância genômica regional, gerar conhecimento "
         "acionável sobre patógenos circulantes e desenvolver uma camada de inteligência "
         "preditiva que amplia a janela de antecipação a surtos. A proposta é explicitamente "
         "complementar — e não concorrente — a sistemas consolidados de monitoramento em tempo "
         "real, como o InfoDengue e o InfoGripe, atuando na etapa anterior de priorização de "
         "ameaças ainda não caracterizadas. Ao produzir ferramentas abertas, reprodutíveis e "
         "adaptadas ao contexto nacional, o projeto reforça a soberania tecnológica em saúde e a "
         "formação de recursos humanos qualificados, alinhando-se à missão institucional da "
         "Fiocruz e às diretrizes de fortalecimento da vigilância genômica pós-COVID-19.")

    # ---------------- 3. Objetivos ----------------
    h1(doc, "3. Objetivos")
    h2(doc, "3.1. Objetivo geral")
    para(doc,
         "Desenvolver e consolidar uma plataforma integrada de bioinformática e inteligência "
         "computacional para caracterização molecular, monitoramento e predição de risco de "
         "vírus emergentes e zoonóticos de importância em saúde pública, com aplicação ao "
         "contexto brasileiro e integração às redes de vigilância do SUS.")
    h2(doc, "3.2. Objetivos específicos")
    bullet(doc, "Estruturar um fluxo automatizado de aquisição e curadoria de genomas virais e "
                "metadados de hospedeiro, taxonomia e ecologia a partir de bases públicas e de "
                "redes de vigilância.")
    bullet(doc, "Caracterizar molecularmente e reconstruir a filogenia e a filodinâmica de vírus "
                "prioritários circulantes no Brasil, com ênfase em arbovírus e vírus respiratórios.")
    bullet(doc, "Desenvolver métodos de aprendizado de máquina auto-supervisionado (redes JEPA) "
                "para aprender representações latentes de genomas virais e priorizar vírus com "
                "potencial de transbordamento zoonótico.", bold_prefix="(*) ")
    bullet(doc, "Integrar os resultados de priorização à vigilância genômica e epidemiológica, "
                "propondo mecanismos de articulação com sistemas como InfoDengue/InfoGripe.")
    bullet(doc, "Formar recursos humanos e disponibilizar código, modelos e documentação de forma "
                "aberta e reprodutível, fortalecendo a capacidade analítica institucional.")
    para(doc,
         "(*) Objetivo específico desenvolvido, no âmbito do Pós-Doutorado Júnior, pelo subprojeto "
         "JEPA-Spillover (candidato Gabriel Bezerra Motta Câmara).",
         size=9.5, italic=True, color=CINZA)

    # ---------------- 4. Metodologia ----------------
    h1(doc, "4. Desenho e abordagem teórico-metodológica")
    para(doc,
         "O projeto adota um desenho modular em quatro eixos interdependentes, combinando "
         "bioinformática genômica clássica, epidemiologia molecular e aprendizado de máquina. A "
         "abordagem prioriza reprodutibilidade (fluxos versionados, ambientes declarados, seeds "
         "fixas) e ciência aberta.")
    h2(doc, "Eixo 1 — Aquisição e curadoria de dados")
    para(doc,
         "Coleta automatizada de sequências e metadados de bases públicas (NCBI Virus/Entrez, "
         "BV-BRC/ViPR, VirusHostDB, NCBI Taxonomy, ICTV) e de dados de acesso controlado (GISAID: "
         "EpiCoV, EpiArbo, EpiNiV), mediante credenciamento institucional. Curadoria com controle "
         "de qualidade (comprimento, bases ambíguas), remoção de duplicatas, padronização "
         "taxonômica e integração de metadados de hospedeiro e interações vírus-hospedeiro "
         "(IntAct/EMBL-EBI).")
    h2(doc, "Eixo 2 — Caracterização molecular, filogenia e filodinâmica")
    para(doc,
         "Montagem e anotação de genomas, alinhamentos, inferência filogenética (máxima "
         "verossimilhança e abordagens bayesianas) e análise filodinâmica para reconstruir "
         "introduções, dispersão e dinâmica temporal de vírus prioritários, com ênfase em "
         "arbovírus e vírus respiratórios circulantes no Brasil. Uso de fluxos padronizados do "
         "tipo Nextstrain para visualização e compartilhamento.")
    h2(doc, "Eixo 3 — Modelagem computacional e IA (subprojeto JEPA-Spillover)")
    para(doc,
         "Desenvolvimento de representações genômicas (k-mers, embeddings) e de uma arquitetura "
         "JEPA genômica que aprende, sem reconstrução literal de nucleotídeos e com baixa "
         "dependência de rótulos, a prever embeddings de regiões mascaradas do genoma. Extensão "
         "vírus-hospedeiro para aprender compatibilidades latentes (incluindo Homo sapiens) e "
         "fine-tuning supervisionado para risco de spillover, validado com particionamento "
         "estratificado por família e holdout de famílias inteiras. Geração de um ranking de "
         "vírus prioritários combinando proximidade latente a zoonóticos conhecidos e "
         "probabilidade supervisionada. Este eixo constitui o subprojeto de PDJ.")
    h2(doc, "Eixo 4 — Integração com a vigilância e disseminação")
    para(doc,
         "Tradução dos resultados para uso em vigilância: painéis interativos, rankings de "
         "priorização e proposta de mecanismos de integração com sistemas de monitoramento "
         "(InfoDengue/InfoGripe) e redes genômicas da Fiocruz. Disseminação via repositórios "
         "públicos, manuscritos, material didático e capacitação de equipes.")

    # ---------------- 5. Articulação ----------------
    h1(doc, "5. Articulação com o subprojeto JEPA-Spillover")
    para(doc,
         "O subprojeto de Pós-Doutorado Júnior JEPA-Spillover materializa o Eixo 3 deste projeto. "
         "Enquanto o projeto guarda-chuva provê o contexto biológico, os dados curados (Eixo 1) e "
         "a caracterização molecular/filodinâmica (Eixo 2), o subprojeto aporta a inovação "
         "metodológica em inteligência artificial: a primeira aplicação, ao domínio de vigilância "
         "genômica de vírus de RNA, das arquiteturas preditivas em espaço latente (JEPA). A "
         "expertise do candidato em aprendizado de máquina e desenvolvimento de software "
         "reprodutível complementa diretamente a competência do grupo em biologia molecular e "
         "epidemiologia genômica, criando sinergia com potencial de publicações conjuntas, "
         "captação de novos financiamentos e consolidação de uma linha de pesquisa em IA aplicada "
         "à saúde no IAM. Os produtos do subprojeto (embeddings, modelos, rankings) retroalimentam "
         "os Eixos 2 e 4, fechando o ciclo entre caracterização, predição e vigilância.")

    # ---------------- 6. Resultados esperados ----------------
    h1(doc, "6. Resultados esperados e impacto")
    bullet(doc, "Banco integrado, curado e documentado de genomas virais e metadados, reutilizável pela instituição.")
    bullet(doc, "Caracterização molecular e filodinâmica de vírus prioritários circulantes no Brasil.")
    bullet(doc, "Modelo JEPA genômico e classificador de risco de spillover, com ranking de vírus prioritários para vigilância.")
    bullet(doc, "Plataforma/painel aberto e reprodutível, com proposta de integração à vigilância (InfoDengue/InfoGripe).")
    bullet(doc, "Publicações científicas, material didático em português e capacitação de recursos humanos em IA aplicada à saúde.")
    bullet(doc, "Fortalecimento da capacidade analítica e da competitividade do IAM/Fiocruz em editais nacionais e internacionais.")

    # ---------------- 7. Infraestrutura ----------------
    h1(doc, "7. Infraestrutura disponível")
    para(doc,
         "O projeto é conduzido no Instituto Aggeu Magalhães (Fiocruz-PE), que dispõe de "
         "infraestrutura de bioinformática e computação científica, incluindo servidores e "
         "recursos de processamento (CPU/GPU) para análises genômicas e treinamento de modelos de "
         "aprendizado profundo, além de acesso às redes e plataformas de vigilância genômica da "
         "Fiocruz. O grupo conta com experiência consolidada em análise de dados de sequenciamento "
         "e epidemiologia molecular, e com credenciamento institucional para bases de acesso "
         "controlado (GISAID). A componente de IA (subprojeto) utiliza estação com GPU e "
         "publicação aberta de código e modelos (GitHub, Hugging Face). [Detalhar equipamentos, "
         "servidores/HPC e espaço físico específicos do NBI/IAM.]")

    # ---------------- 8. Cronograma ----------------
    h1(doc, "8. Cronograma (36 meses)")
    make_table(
        doc,
        ["Fase", "Atividades", "Período", "Eixo"],
        [
            ["1", "Estruturação do fluxo de dados; curadoria e banco integrado", "Meses 1–6", "1"],
            ["2", "Caracterização molecular, filogenia e filodinâmica de vírus prioritários", "Meses 4–18", "2"],
            ["3", "Representações, JEPA genômica e extensão vírus-hospedeiro (PDJ)", "Meses 7–18", "3"],
            ["4", "Fine-tuning de risco, validação e ranking de priorização (PDJ)", "Meses 13–24", "3"],
            ["5", "Integração com vigilância; painéis e proposta de articulação", "Meses 19–30", "4"],
            ["6", "Disseminação: manuscritos, material didático, capacitação, relatório", "Meses 25–36", "4"],
        ],
        widths=[1.6, 9.4, 3.4, 1.6],
        font_size=9,
    )
    para(doc,
         "O subprojeto de Pós-Doutorado Júnior (JEPA-Spillover) concentra-se nas Fases 3 e 4 "
         "(Eixo 3), com vigência inicial de 12 meses, renovável, e é plenamente exequível dentro "
         "desse período — parte substancial da infraestrutura de software já se encontra "
         "implementada em fase de prova de conceito.",
         size=10)

    # ---------------- 9. Equipe ----------------
    h1(doc, "9. Equipe e colaborações")
    bullet(doc, "Túlio Campos [confirmar] — coordenação, bioinformática e epidemiologia molecular (IAM/Fiocruz).", bold_prefix="Supervisor: ")
    bullet(doc, "Gabriel Bezerra Motta Câmara — modelagem computacional e IA (subprojeto JEPA-Spillover).", bold_prefix="Bolsista PDJ: ")
    bullet(doc, "[Demais membros do Núcleo de Bioinformática e colaboradores do IAM].", bold_prefix="Equipe: ")
    bullet(doc, "[Redes de vigilância genômica da Fiocruz; colaborações nacionais e internacionais].", bold_prefix="Colaborações: ")

    # ---------------- 10. Referências ----------------
    h1(doc, "10. Referências selecionadas")
    refs = [
        "LeCun, Y. A Path Towards Autonomous Machine Intelligence. Open Review, 2022.",
        "Assran, M. et al. Self-Supervised Learning from Images with a Joint-Embedding Predictive Architecture (I-JEPA). CVPR, 2023.",
        "Carroll, D. et al. The Global Virome Project. Science, 2018.",
        "Mollentze, N.; Streicker, D. G. et al. Identifying and prioritizing potential human-infecting viruses from their genome sequences. PLoS Biology, 2021.",
        "Hadfield, J. et al. Nextstrain: real-time tracking of pathogen evolution. Bioinformatics, 2018.",
        "Faria, N. R. et al. Genomic and epidemiological monitoring of arboviruses (Zika/dengue) in Brazil. Science, 2016/2017.",
        "Codeço, C. T. et al. InfoDengue: sistema de alerta de arboviroses. Fiocruz.",
        "Mihara, T. et al. Virus-Host Database (VHDB). Viruses, 2016.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{i}] {ref}")
        r.font.size = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    footer = sec.footer
    fp = footer.paragraphs[0]
    fr = fp.add_run("Projeto do(a) Supervisor(a) — PDJ · Instituto Aggeu Magalhães / Fiocruz PE · 2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = CINZA
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
