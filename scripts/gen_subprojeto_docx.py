#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=============================================================================
JEPA-Spillover — gerador do Subprojeto do Candidato (PDJ)
=============================================================================
Projeto : JEPA-Spillover (PDJ / IAM — Fiocruz PE)
Módulo  : scripts/gen_subprojeto_docx.py

Propósito
---------
Gera o documento Word do **Subprojeto do Candidato** (item 4.4.f do Edital
PDJ/Fiocruz VPPCB 2026), alinhado aos critérios do Anexo IV.

Supervisor / lab atuais
-----------------------
Marcelo Henrique Santos Paiva — Laboratório de Entomologia (IAM).

Saídas
------
- ``submissao_pdj/Subprojeto_JEPA-Spillover_PDJ.docx``
- Converter para PDF: ``libreoffice --headless --convert-to pdf …``

Uso
---
    python scripts/gen_subprojeto_docx.py

Campos em aberto
----------------
Infraestrutura do laboratório e aderência da trajetória (placeholders ``[ ]``).
=============================================================================
"""

from __future__ import annotations

import logging

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
log = logging.getLogger("scripts.gen_subprojeto_docx")

OUT = ROOT / "submissao_pdj" / "Subprojeto_JEPA-Spillover_PDJ.docx"

# ---- Paleta / constantes ----------------------------------------------------
AZUL = RGBColor(0x1F, 0x39, 0x64)
CINZA = RGBColor(0x55, 0x55, 0x55)
VERDE = RGBColor(0x1B, 0x5E, 0x20)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    tc_pr.append(shd)


def style_header_row(row, hex_bg="1F3964") -> None:
    for cell in row.cells:
        set_cell_bg(cell, hex_bg)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(14)
    p.space_before = Pt(10)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(12)
    return p


def para(doc, text, size=11, justify=True, italic=False, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def make_table(doc, headers, rows, widths=None, header_bg="1F3964", font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].add_run(htext)
    style_header_row(t.rows[0], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pr = cells[i].paragraphs[0]
            run = pr.add_run(str(val))
            run.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


# =============================================================================
def build():
    """Monta e salva o Subprojeto do Candidato (docx) conforme o edital PDJ."""
    doc = Document()

    # Estilo base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # ---- CAPA ---------------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FUNDAÇÃO OSWALDO CRUZ — FIOCRUZ")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = AZUL

    for txt, sz in [
        ("Vice-Presidência de Pesquisa e Coleções Biológicas (VPPCB)", 11),
        ("Instituto Aggeu Magalhães — Fiocruz Pernambuco", 11),
        ("Edital de Pós-Doutorado Júnior (PDJ) — Vigência 2026–2027", 11),
    ]:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = q.add_run(txt)
        rr.font.size = Pt(sz)
        rr.font.color.rgb = CINZA

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SUBPROJETO DO CANDIDATO")
    r.font.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "JEPA-Spillover: aprendizado preditivo em espaço latente para "
        "vigilância genômica de vírus com potencial zoonótico"
    )
    r.font.bold = True
    r.font.size = Pt(14)
    r.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Bloco de identificação
    ident = [
        ("Candidato", "Gabriel Bezerra Motta Câmara"),
        ("Supervisor(a)", "Marcelo Henrique Santos Paiva — servidor(a) ativo(a) da Fiocruz"),
        ("Projeto do(a) supervisor(a)", "GenVig-IA — Vigilância Genômica Integrada e Inteligência Computacional para vírus emergentes e zoonóticos (Eixo 3)"),
        ("Unidade", "Instituto Aggeu Magalhães (IAM) — Fiocruz Pernambuco"),
        ("Laboratório / Grupo", "Laboratório de Entomologia"),
        ("Modalidade", "Bolsa Nova"),
        ("Vigência proposta", "12 meses (mínimo do edital) — prorrogável conforme normas"),
        ("Grande área", "Ciências da Saúde / Bioinformática / Inteligência Artificial aplicada à Saúde"),
        ("Linha de pesquisa", "Vigilância genômica, saúde pública e predição de risco zoonótico"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in ident:
        cells = t.add_row().cells
        rk = cells[0].paragraphs[0].add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)
        set_cell_bg(cells[0], "EDF1F7")
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(10)
        cells[0].width = Cm(5.2)
        cells[1].width = Cm(11.3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Recife — PE · 2026")
    r.font.size = Pt(11)
    r.font.color.rgb = CINZA

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = nota.add_run(
        "Campos entre colchetes [ ] restantes devem ser preenchidos antes da submissão "
        "no sistema Fomento à Pesquisa on-line (ex.: infraestrutura do laboratório — em aberto por enquanto)."
    )
    rn.italic = True
    rn.font.size = Pt(8.5)
    rn.font.color.rgb = CINZA

    doc.add_page_break()

    # ---- 1. RESUMO ----------------------------------------------------------
    h1(doc, "1. Resumo")
    para(doc,
         "Eventos de transbordamento zoonótico (spillover) estão entre as maiores ameaças à saúde "
         "pública global, como evidenciado pela pandemia de COVID-19. O volume de genomas virais "
         "cresce rapidamente por vigilância genômica, metagenômica ambiental e estudos de viroma, "
         "mas a maioria dos vírus recém-identificados não possui rótulos confiáveis de risco "
         "zoonótico — o que limita fortemente os métodos supervisionados tradicionais. Este "
         "subprojeto desenvolve uma abordagem auto-supervisionada baseada em redes JEPA (Joint "
         "Embedding Predictive Architecture) para aprender representações latentes de genomas "
         "virais, hospedeiros e contextos ecológicos, com o objetivo de priorizar vírus com maior "
         "potencial de spillover antes da ocorrência de surtos. A proposta é complementar — não "
         "concorrente — aos sistemas nacionais de vigilância já consolidados (InfoDengue, InfoGripe, "
         "SIVEP-Gripe), oferecendo uma camada de inteligência genômica que amplia a janela de "
         "antecipação. Todo o pipeline é aberto, reprodutível e já parcialmente implementado, com "
         "resultados preliminares promissores de baselines e infraestrutura computacional funcional.")

    # ---- 2. INTRODUÇÃO E JUSTIFICATIVA -------------------------------------
    h1(doc, "2. Introdução e justificativa")
    para(doc,
         "A predição proativa de risco zoonótico enfrenta limitações estruturais que a abordagem "
         "JEPA busca superar. Modelos supervisionados dependem de rótulos 'zoonótico/não-zoonótico' "
         "escassos, incompletos e enviesados — um vírus classificado como 'não-zoonótico' pode "
         "simplesmente não ter sido suficientemente estudado. Sistemas como o InfoDengue e o "
         "InfoGripe são altamente eficazes no monitoramento em tempo real de vírus já conhecidos e "
         "circulantes, mas não foram projetados para identificar vírus desconhecidos com potencial "
         "zoonótico antes de um surto, pois métodos estatísticos tradicionais dependem de dados "
         "históricos de transmissão — inexistentes para vírus emergentes recém-descobertos.")
    para(doc,
         "Métodos de aprendizado profundo baseados em reconstrução (autoencoders, modelos de "
         "linguagem nucleotídeo a nucleotídeo) gastam capacidade modelando ruído de sequência de "
         "baixo nível em vez de semântica funcional. A JEPA, ao contrário, prevê representações "
         "latentes — não nucleotídeos — capturando relações biológicas relevantes com maior "
         "eficiência e baixa dependência de rótulos, o que a torna especialmente adequada para "
         "vírus pouco caracterizados.")

    h2(doc, "2.1. Por que JEPA?")
    make_table(
        doc,
        ["Abordagem", "Limitação no contexto de spillover"],
        [
            ["Supervisionada (rótulos zoonótico/não)",
             "Rótulos escassos, incompletos e enviesados; 'não zoonótico' pode significar apenas 'não estudado'."],
            ["Generativa / reconstrução (autoencoders, MLM nucleotídeo a nucleotídeo)",
             "Gasta capacidade modelando ruído de baixo nível em vez de semântica funcional."],
            ["JEPA (predição em espaço latente)",
             "Aprende relações latentes entre partes dos dados sem reconstrução literal e com baixa dependência de rótulos — ideal para vírus pouco caracterizados."],
        ],
        widths=[6.0, 10.5],
    )
    doc.add_paragraph()
    para(doc,
         "A arquitetura JEPA, proposta por Yann LeCun (Meta AI / NYU) como princípio fundador dos "
         "'modelos de mundo' (World Models), representa hoje a fronteira do aprendizado de máquina "
         "auto-supervisionado. Nos últimos anos, uma série de implementações estabeleceu recordes em "
         "benchmarks de aprendizado sem rótulos:")
    bullet(doc, "primeira implementação pública; treinada sem rótulos, supera MAE e DINO v2 em classificação linear e few-shot no ImageNet ao prever representações latentes em vez de pixels.", bold_prefix="I-JEPA (Image, 2023): ")
    bullet(doc, "extensão para vídeo; aprende representações espaço-temporais e supera modelos de difusão em compreensão de ação com ordens de magnitude menos parâmetros, demonstrando escalabilidade para sequências.", bold_prefix="V-JEPA (Video, 2024): ")
    bullet(doc, "integra predição latente com aprendizado contrastivo multimodal — diretamente relevante para integrar sequências genômicas a metadados de hospedeiro e contexto epidemiológico.", bold_prefix="MC-JEPA (Multimodal, 2023): ")
    bullet(doc, "aprendem representações em múltiplas escalas; para genômica, capturam simultaneamente motivos curtos (códons, sítios de ligação) e a estrutura global do genoma.", bold_prefix="JEPA hierárquicas (2024–2025): ")
    para(doc,
         "Três características são especialmente vantajosas para sequências virais: "
         "(i) predição no espaço latente, não no espaço de entrada, forçando entendimento funcional "
         "em vez de memorização de composição de bases e ruído de sequenciamento; (ii) "
         "codificador-alvo por média móvel exponencial (EMA, stop-gradient), que gera sinal de treino "
         "estável e progressivamente mais rico, análogo à emergência de representações biológicas por "
         "pressão seletiva gradual; e (iii) dispensa de pares negativos, difíceis de definir em dados "
         "genômicos.")
    para(doc,
         "O JEPA-Spillover é — até onde conhecemos — a primeira aplicação da arquitetura JEPA ao "
         "domínio de vigilância genômica de vírus de RNA. Trabalhos anteriores em genômica "
         "auto-supervisionada (DNABERT, Nucleotide Transformer, HyenaDNA) usam variantes de BERT "
         "mascarado ou modelos autoregressivos que operam no espaço de tokens. Este posicionamento "
         "na fronteira metodológica constitui diferencial concreto para captação de financiamento, "
         "colaborações internacionais e formação de recursos humanos especializados.",
         )

    # ---- 3. ARTICULAÇÃO COM O PROJETO DO SUPERVISOR ------------------------
    h1(doc, "3. Articulação com o projeto do(a) supervisor(a)")
    para(doc,
         "Este subprojeto vincula-se diretamente ao projeto de pesquisa do(a) supervisor(a) "
         "\u201cGenVig-IA — Vigilância Genômica Integrada e Inteligência Computacional para "
         "Caracterização Molecular e Predição de Risco de Vírus Emergentes e Zoonóticos de "
         "Importância em Saúde Pública no Brasil\u201d, do qual constitui a componente de "
         "inteligência computacional (Eixo 3). Enquanto o projeto guarda-chuva abrange a aquisição "
         "e curadoria de dados genômicos (Eixo 1) e a caracterização molecular, filogenética e "
         "filodinâmica de vírus prioritários circulantes no Brasil (Eixo 2), o subprojeto aporta "
         "métodos de aprendizado de máquina de última geração (JEPA/Transformers) para "
         "priorização proativa de vírus com potencial zoonótico, agregando capacidade analítica "
         "reprodutível e escalável ao grupo. Os produtos do subprojeto (embeddings, modelos e "
         "rankings de risco) retroalimentam os eixos de caracterização e de integração com a "
         "vigilância (Eixo 4). A convergência entre a expertise em entomologia e saúde pública do "
         "Laboratório de Entomologia e a competência computacional do candidato cria sinergia com potencial de "
         "gerar publicações conjuntas, novos editais e colaborações.")

    # ---- 4. OBJETIVO GERAL --------------------------------------------------
    h1(doc, "4. Objetivo geral")
    para(doc,
         "Desenvolver uma abordagem baseada em redes JEPA para aprender representações latentes de "
         "genomas virais, hospedeiros e contextos ecológicos, apoiando a predição e a priorização "
         "de vírus com potencial de spillover, de forma aberta, reprodutível e integrável às redes "
         "de vigilância em saúde pública do Brasil.")

    # ---- 5. OBJETIVOS ESPECÍFICOS ------------------------------------------
    h1(doc, "5. Objetivos específicos")
    for oe in [
        "Coletar e curar genomas virais públicos e integrar metadados de hospedeiros, taxonomia e ecologia em um banco padronizado e documentado.",
        "Construir representações genômicas (k-mers e embeddings baseados em Transformer) e estabelecer baselines comparativos.",
        "Implementar e treinar a JEPA genômica (contexto → alvo latente) para gerar embeddings virais.",
        "Estender a arquitetura para o par vírus–hospedeiro e, de forma exploratória, integrar contexto ecológico em um espaço latente compartilhado.",
        "Realizar fine-tuning supervisionado para risco de spillover e validar com particionamento estratificado por família e holdout de famílias inteiras.",
        "Aplicar interpretabilidade (SHAP), ablação e visualização (UMAP/t-SNE); produzir um ranking de vírus prioritários para vigilância.",
        "Disponibilizar código, modelos e documentação em repositório público reprodutível e promover capacitação da equipe.",
    ]:
        bullet(doc, oe)

    # ---- 6. ABORDAGEM TEÓRICO-METODOLÓGICA E MÉTODOS -----------------------
    h1(doc, "6. Abordagem teórico-metodológica e métodos")

    h2(doc, "6.1. Fontes de dados e curadoria")
    para(doc,
         "Scripts automatizados baixam sequências de bases públicas — NCBI Virus (via Entrez/"
         "Biopython), VirusHostDB (relações vírus–hospedeiro), NCBI Taxonomy (padronização), IntAct/"
         "VirHostNet (interações moleculares) e GISAID (credencial institucional; não redistribuído). "
         "A curadoria aplica controle de qualidade (comprimento mínimo, fração de bases ambíguas), "
         "remoção de duplicatas e padronização taxonômica, resultando em um dataset em formato "
         "Parquet com sequência, metadados de hospedeiro, taxonomia e rótulos de risco.")

    h2(doc, "6.2. Representações e baselines")
    bullet(doc, "frequências de k-mers reduzidas por PCA (baseline rápido, sem GPU).", bold_prefix="k-mer + PCA: ")
    bullet(doc, "embeddings por codificadores de DNA pré-treinados (ex.: Nucleotide Transformer).", bold_prefix="Transformer: ")
    bullet(doc, "k-mers + Regressão Logística/Random Forest, LSTM supervisionada e Transformer supervisionado.", bold_prefix="Baselines: ")

    h2(doc, "6.3. Arquitetura JEPA genômica")
    bullet(doc, "codifica um bloco observado do genoma.", bold_prefix="Codificador de contexto f(θ) (Transformer): ")
    bullet(doc, "gera o embedding de blocos mascarados (stop-gradient).", bold_prefix="Codificador-alvo por EMA: ")
    bullet(doc, "prevê o embedding-alvo a partir do contexto + posição.", bold_prefix="Preditor g(φ): ")
    para(doc,
         "A perda é calculada no espaço latente (Smooth L1), não sobre nucleotídeos, incentivando o "
         "aprendizado de relações funcionais entre regiões do genoma.")

    h2(doc, "6.4. Extensão vírus–hospedeiro e fine-tuning")
    para(doc,
         "O embedding viral passa a prever o embedding do hospedeiro associado, aprendendo "
         "compatibilidades latentes vírus–hospedeiro (incluindo Homo sapiens); uma versão "
         "exploratória integra embeddings ecológicos. Em seguida, os embeddings JEPA alimentam um "
         "classificador leve treinado com rótulos derivados de histórico de infecção humana, "
         "amplitude de hospedeiros, surtos e transmissão interespécies.")

    h2(doc, "6.5. Avaliação, interpretabilidade e ranking")
    bullet(doc, "AUROC, AUPRC, F1, precisão, recall, especificidade e Brier.", bold_prefix="Métricas: ")
    bullet(doc, "validação cruzada estratificada por família (5 folds) + holdout de famílias inteiras para simular vírus emergentes.", bold_prefix="Validação: ")
    bullet(doc, "SHAP + ablação (contribuição de regiões genômicas e variáveis) e projeções UMAP/t-SNE por família, hospedeiro e risco.", bold_prefix="Interpretabilidade: ")
    bullet(doc, "score de prioridade combinando proximidade latente a zoonóticos conhecidos e probabilidade supervisionada — vírus próximos a Ebola/SARS sobem no ranking automaticamente.", bold_prefix="Ranking: ")

    # ---- 7. ATIVIDADES E METAS ---------------------------------------------
    h1(doc, "7. Atividades e metas")
    para(doc,
         "As metas abaixo articulam-se diretamente aos objetivos específicos (Seção 5) e às "
         "atividades do cronograma (Seção 8), garantindo coerência entre objetivos, atividades e "
         "produtos esperados.")
    make_table(
        doc,
        ["Meta", "Atividades", "Entregável principal"],
        [
            ["1 — Base integrada",
             "Levantamento de bases, coleta, QC, dedup e padronização taxonômica",
             "Banco tabular/relacional + documentação de curadoria"],
            ["2 — Representações + baselines",
             "k-mers, embeddings e implementação de baselines",
             "Scripts de pré-processamento, matriz de atributos, baselines"],
            ["3 — JEPA genômica",
             "Implementação, treinamento e avaliação de embeddings",
             "Modelo funcional + checkpoints + avaliação de embeddings"],
            ["4 — Extensão vírus-hospedeiro",
             "Modelo vírus-hospedeiro + integração ecológica exploratória",
             "Modelo JEPA vírus-hospedeiro + espaço latente compartilhado"],
            ["5 — Fine-tuning + validação",
             "Classificador de risco, validação, comparação e ranking",
             "Métricas, comparação e ranking preliminar de vírus prioritários"],
            ["6 — Disseminação",
             "Repositório, manuscrito, material didático e relatório final",
             "Repositório documentado + manuscrito + apresentação"],
        ],
        widths=[3.6, 7.0, 5.9],
    )

    # ---- 8. CRONOGRAMA ------------------------------------------------------
    h1(doc, "8. Cronograma de execução (12 meses)")
    para(doc,
         "O cronograma é exequível dentro da vigência mínima de 12 meses do edital (item de "
         "avaliação 'Exequibilidade'), com marcos claros e produtos verificáveis a cada trimestre. "
         "Parte substancial da infraestrutura de software já está implementada (ver Seção 9), o que "
         "reduz o risco de execução.")
    make_table(
        doc,
        ["Mês", "Atividades principais", "Produtos esperados", "Meta"],
        [
            ["1", "Levantamento de bases, critérios de inclusão, desenho do pipeline", "Plano técnico + lista de bases", "1"],
            ["2", "Coleta e curadoria inicial de genomas e metadados", "Base preliminar curada", "1"],
            ["3", "Padronização taxonômica, dedup, estruturação", "Banco integrado v1", "1"],
            ["4", "k-mers, embeddings genômicos, baselines", "Matrizes de atributos + baselines", "2"],
            ["5", "Implementação da JEPA genômica", "Protótipo funcional", "3"],
            ["6", "Treinamento e ajuste da JEPA genômica", "Embeddings virais preliminares", "3"],
            ["7", "Extensão vírus-hospedeiro", "Modelo JEPA vírus-hospedeiro", "4"],
            ["8", "Integração ecológica/interacional exploratória", "Espaço latente multimodal preliminar", "4"],
            ["9", "Fine-tuning supervisionado para risco", "Modelo preditivo + métricas iniciais", "5"],
            ["10", "Validação e comparação (LSTM/Transformer/k-mers)", "Relatório comparativo", "5"],
            ["11", "Interpretabilidade, ablação, UMAP/t-SNE, ranking", "Figuras, rankings, interpretação", "5"],
            ["12", "Repositório, manuscrito, relatório final", "Manuscrito + repositório + relatório", "6"],
        ],
        widths=[1.3, 7.2, 6.4, 1.1],
        font_size=9,
    )
    doc.add_paragraph()
    para(doc, "Marcos (milestones):", bold=True)
    bullet(doc, "Banco integrado v1 disponível e documentado.", bold_prefix="M3 — ")
    bullet(doc, "JEPA genômica treinada com embeddings preliminares avaliados.", bold_prefix="M6 — ")
    bullet(doc, "Primeiro modelo de risco de spillover com métricas.", bold_prefix="M9 — ")
    bullet(doc, "Repositório público + manuscrito em preparação/submissão.", bold_prefix="M12 — ")

    # ---- 9. RESULTADOS PRELIMINARES ----------------------------------------
    h1(doc, "9. Resultados preliminares")
    para(doc,
         "O subprojeto já conta com implementação avançada e resultados preliminares que demonstram "
         "viabilidade técnica:")
    bullet(doc, "~59 mil sequências coletadas (13 famílias virais do NCBI + GISAID) e 26 mil após curadoria, com balanceamento por família (2.000/família).", bold_prefix="Dados: ")
    bullet(doc, "pipeline completo em Python (coleta, curadoria, features, treino, avaliação, dashboard) versionado e documentado.", bold_prefix="Software: ")
    bullet(doc, "pré-treino JEPA em GPU (NVIDIA RTX 3050) com checkpointing por época e retomada automática.", bold_prefix="Treino: ")
    para(doc,
         "Benchmark de representações k-mer (3.000 sequências, 128 componentes PCA) — seleção "
         "baseada em desempenho e custo de memória:")
    make_table(
        doc,
        ["k", "Features", "Var. explicada", "AUROC (5-fold)", "RAM"],
        [
            ["3", "64", "100%", "0,9948", "2 MB"],
            ["4 (selecionado)", "256", "98,2%", "0,9961", "8 MB"],
            ["5", "1.024", "87,2%", "0,9961", "36 MB"],
            ["6", "4.096", "75,4%", "0,9947", "115 MB"],
            ["7", "16.384", "68,3%", "0,9934", "451 MB"],
        ],
        widths=[3.4, 2.6, 3.4, 3.6, 2.6],
    )
    doc.add_paragraph()
    para(doc,
         "O valor k = 4 foi selecionado por apresentar o melhor AUROC (empatado com k = 5) usando "
         "56× menos memória que k = 7. Notavelmente, k = 7 apresenta o pior AUROC — vocabulários "
         "grandes introduzem ruído em datasets deste porte. Estes baselines fortes estabelecem o "
         "piso de desempenho a ser superado pelos embeddings JEPA na predição de risco entre "
         "famílias virais (cenário mais realista para vírus emergentes).",
         italic=False)
    para(doc,
         "Recursos públicos do projeto: código em github.com/gbmotta/jepa-spillover; modelos e "
         "dados em huggingface.co/gbmotta; dashboard interativo em "
         "huggingface.co/spaces/gbmotta/jepa-spillover-dashboard.")

    # ---- 10. INFRAESTRUTURA -------------------------------------------------
    h1(doc, "10. Infraestrutura disponível")
    para(doc,
         "O subprojeto é predominantemente computacional, o que reduz custos e dependências "
         "laboratoriais. Descrever a infraestrutura disponível no Laboratório de Entomologia "
         "e no IAM/Fiocruz para a realização das atividades previstas:")
    para(doc,
         "[Computação local/GPU, servidores/HPC, ambiente de software, acesso a bases de dados, "
         "publicação/reprodutibilidade, espaço físico e demais recursos institucionais — "
         "a preencher com o(a) supervisor(a).]",
         italic=True, color=CINZA)

    # ---- 11. CONTRIBUIÇÃO PARA O GRUPO -------------------------------------
    h1(doc, "11. Contribuição do subprojeto para o grupo de pesquisa / laboratório")
    para(doc,
         "Além da entrega de modelos e código, o subprojeto prevê a transferência de conhecimento "
         "em inteligência artificial, deep learning e aprendizado auto-supervisionado para "
         "estudantes e membros da equipe. Esse componente formativo é um produto duradouro: mesmo "
         "que o modelo JEPA seja refinado em etapas futuras, as competências desenvolvidas "
         "permanecem na instituição e podem ser aplicadas a outros contextos. Como diferenciais "
         "concretos, o subprojeto pode incluir:")
    bullet(doc, "tutoriais e material didático em português sobre JEPA aplicada à bioinformática viral;")
    bullet(doc, "workshops internos sobre aprendizado auto-supervisionado, embeddings e análise de genomas com IA;")
    bullet(doc, "capacitação de colaboradores de outros grupos do IAM/Fiocruz interessados em IA aplicada à saúde;")
    bullet(doc, "documentação e scripts reutilizáveis como base para projetos futuros do laboratório.")
    para(doc,
         "Ao aplicar arquiteturas de aprendizado profundo de última geração a problemas de "
         "vigilância genômica, o subprojeto posiciona o laboratório e o Instituto Aggeu Magalhães na "
         "fronteira da bioinformática aplicada à saúde pública, fortalecendo a competitividade em "
         "editais nacionais e internacionais e abrindo portas para colaborações científicas.")

    # ---- 12. CONTRIBUIÇÃO DA FORMAÇÃO DO CANDIDATO -------------------------
    h1(doc, "12. Contribuição da formação do candidato para o projeto do(a) supervisor(a)")
    para(doc,
         "A formação e a expertise técnico-científica do candidato — em ciência de dados, "
         "aprendizado de máquina e desenvolvimento de software reprodutível — complementam "
         "diretamente as competências em entomologia e saúde pública do Laboratório de Entomologia. O candidato aporta a "
         "capacidade de implementar, treinar e avaliar modelos de IA de fronteira (JEPA/"
         "Transformers), estruturar pipelines de dados escaláveis e publicar resultados de forma "
         "aberta e reprodutível. Essa integração acelera a execução do projeto do(a) supervisor(a), "
         "amplia seu alcance metodológico e cria condições para produção científica conjunta de alto "
         "impacto, submissão a novos editais e formação de recursos humanos qualificados no IAM/"
         "Fiocruz. [Complementar com aderência específica entre a trajetória do candidato — "
         "titulação, publicações e experiência — e as necessidades do projeto do(a) supervisor(a).]")

    # ---- 13. ALINHAMENTO COM O SUS -----------------------------------------
    h1(doc, "13. Relevância para o SUS e alinhamento com a vigilância nacional")
    para(doc,
         "A proposta é complementar — não concorrente — aos sistemas de vigilância consolidados no "
         "Brasil. O InfoDengue e o InfoGripe são plataformas de monitoramento epidemiológico em "
         "tempo real amplamente usadas pelo Ministério da Saúde e por secretarias estaduais, com "
         "foco em vírus já conhecidos e circulantes. O JEPA-Spillover atua na etapa anterior: "
         "identificar, entre vírus pouco caracterizados de vigilância ambiental e metagenômica, "
         "aqueles com maior potencial de se tornarem futuros alvos desses sistemas. Como "
         "colaboração futura explícita, vírus priorizados pelo ranking JEPA poderiam ser "
         "incorporados como alertas precoces ao InfoDengue/InfoGripe.")
    para(doc,
         "O subprojeto alinha-se às diretrizes do Ministério da Saúde e da Fiocruz para "
         "fortalecimento da vigilância genômica no contexto pós-COVID-19, e contribui para a "
         "soberania nacional em tecnologias críticas de saúde ao desenvolver soluções de IA abertas, "
         "reprodutíveis e adaptadas ao contexto brasileiro, com potencial de integração a redes como "
         "SIVEP-Gripe, CIEVS e o Sistema de Alerta e Resposta de Emergências.")

    # ---- 14. RISCOS E MITIGAÇÃO --------------------------------------------
    h1(doc, "14. Riscos e mitigação")
    make_table(
        doc,
        ["Risco", "Mitigação"],
        [
            ["Rótulos escassos/enviesados", "Núcleo auto-supervisionado (JEPA); validação entre famílias inteiras."],
            ["Dados ecológicos incompletos", "Integração incremental; componente multimodal tratado como exploratório."],
            ["Custo computacional (Transformers)", "Baseline k-mer/CNN; uso de GPU quando disponível; blocos menores e checkpointing."],
            ["Vazamento de informação na avaliação", "Particionamento estratificado + holdout de famílias inteiras."],
        ],
        widths=[6.0, 10.5],
    )

    # ---- 15. CIÊNCIA ABERTA -------------------------------------------------
    h1(doc, "15. Reprodutibilidade e ciência aberta")
    bullet(doc, "configuração única e versionada (config.yaml) e seed fixa;")
    bullet(doc, "ambiente declarado (environment.yml / requirements.txt);")
    bullet(doc, "acompanhamento de experimentos e publicação de código, modelos e documentação em GitHub e Hugging Face Hub;")
    bullet(doc, "depósito da produção textual no Repositório Institucional ARCA, em conformidade com a Política Institucional de Acesso Aberto da Fiocruz.")

    # ---- 16. ÉTICA ----------------------------------------------------------
    h1(doc, "16. Considerações éticas e de biossegurança")
    para(doc,
         "O subprojeto utiliza exclusivamente sequências genômicas e metadados de bases de dados "
         "públicas e de acesso controlado (GISAID, mediante credenciamento institucional), sem "
         "coleta de material biológico novo nem dados de pacientes identificáveis, o que, em "
         "princípio, dispensa submissão ao CEP/CEUA. O uso de dados GISAID observa integralmente os "
         "termos de uso da plataforma, e as sequências não são redistribuídas. Eventuais "
         "autorizações legais aplicáveis (CEP, CEUA, SISBIO, CI-BIO, SISGEN) serão anexadas conforme "
         "o item 4.6 do edital, quando pertinentes ao projeto do(a) supervisor(a).")

    # ---- 17. AVISO / DISCLAIMER --------------------------------------------
    h1(doc, "17. Aviso")
    para(doc,
         "Trata-se de ferramenta de apoio à priorização para vigilância genômica, em estágio de "
         "prova de conceito. Não substitui investigação laboratorial, sistemas de vigilância "
         "epidemiológica estabelecidos nem decisões de saúde pública. Resultados in silico exigem "
         "validação experimental.",
         italic=True)

    # ---- 18. REFERÊNCIAS ----------------------------------------------------
    h1(doc, "18. Referências selecionadas")
    refs = [
        "Assran, M. et al. Self-Supervised Learning from Images with a Joint-Embedding Predictive Architecture (I-JEPA). CVPR, 2023.",
        "Bardes, A. et al. V-JEPA: Video Joint-Embedding Predictive Architecture. Meta AI, 2024.",
        "LeCun, Y. A Path Towards Autonomous Machine Intelligence. Open Review, 2022.",
        "Ji, Y. et al. DNABERT: pre-trained Bidirectional Encoder Representations from Transformers model for DNA-language in genome. Bioinformatics, 2021.",
        "Dalla-Torre, H. et al. The Nucleotide Transformer: Building and Evaluating Robust Foundation Models for Human Genomics. bioRxiv, 2023.",
        "Mizumoto, K.; Chowell, G. et al. Sistemas de vigilância e modelagem epidemiológica (InfoDengue/InfoGripe). Fiocruz.",
        "Mihara, T. et al. Virus-Host Database (VHDB). Viruses, 2016.",
        "ICTV. Master Species List (MSL). International Committee on Taxonomy of Viruses.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{i}] {ref}")
        r.font.size = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Rodapé simples
    footer = sec.footer
    fp = footer.paragraphs[0]
    fr = fp.add_run("Subprojeto PDJ — JEPA-Spillover · Instituto Aggeu Magalhães / Fiocruz PE · 2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = CINZA
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    log.info("Documento gerado: %s", OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(name)s | %(message)s")
    build()
